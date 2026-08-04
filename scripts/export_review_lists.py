#!/usr/bin/env python3
"""
scripts/export_review_lists.py
Export BioProject review lists as Word documents, one per Sankey category.

8 output files in scripts/review_lists/:
  {setting}_{treatment}.docx   (field/lab × single/host_study/abiotic_stress/other)

Each doc lists one BioProject per block with: accession, title, PMID/publication,
LLM classification + rationale, co-infection summary, named pathogen/host.

Run from crypt/:  python scripts/export_review_lists.py
"""

import csv
import sys
from collections import defaultdict
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("pip install python-docx")

LLM_TSV  = Path("output/05_llm_classify/data/bioproject_llm.tsv")
KW_TSV   = Path("output/04_filter_kw/data/biosample_kw.tsv")
OUT_DIR  = Path("scripts/review_lists")
OUT_DIR.mkdir(exist_ok=True)

TREAT_ORDER  = ["single", "host_study", "abiotic_stress", "other"]
TREAT_LABELS = {
    "single":         "Single-pathogen",
    "host_study":     "Host study",
    "abiotic_stress": "Abiotic stress",
    "other":          "Other / unclear",
}

# ── Load LLM classifications ───────────────────────────────────────────────────

llm = {}
with open(LLM_TSV) as f:
    for r in csv.DictReader(f, delimiter="\t"):
        llm[r["BioProject"]] = r

def llm_setting(bp):
    s = llm.get(bp, {}).get("llm_study_setting", "")
    return "field" if s == "field" else "lab"

def llm_treat(bp):
    t = llm.get(bp, {}).get("llm_treatment", "unclear")
    return t if t in ("single", "host_study", "abiotic_stress") else "other"

# ── Aggregate per BioProject from biosample_kw ────────────────────────────────
# Each BioProject may have multiple BioSamples; collect summary stats.

bp_data = {}  # bp → dict with meta + aggregate stats

with open(KW_TSV) as f:
    for r in csv.DictReader(f, delimiter="\t"):
        bp = r["BioProject"]
        if bp not in bp_data:
            bp_data[bp] = {
                "BioProject":          bp,
                "title":               r.get("title", "").strip(),
                "abstract":            r.get("abstract", "").strip(),
                "primary_pmid":        r.get("primary_pmid", "").strip(),
                "primary_pub_date":    r.get("primary_pub_date", "").strip(),
                "primary_publication": r.get("primary_publication", "").strip(),
                "bp_submission_date":  r.get("bp_submission_date", "").strip(),
                "n_biosamples":        0,
                "n_coinf":             0,
                "n_hc":                0,
                "pathogens":           set(),
                "named_host":          set(),
                "named_pathogen":      set(),
            }
        d = bp_data[bp]
        d["n_biosamples"] += 1
        flag = r.get("co_infection_flag", "single")
        if flag != "single":
            d["n_coinf"] += 1
        if flag != "single" and r.get("same_genus_secondary", "True") == "False":
            d["n_hc"] += 1
        for p in r.get("stat_pathogens", "").split(";"):
            p = p.split(":")[0].strip()
            if p:
                d["pathogens"].add(p)
        nh = r.get("named_host", "").strip()
        if nh:
            d["named_host"].add(nh)

# Merge LLM named pathogen/host
for bp, d in bp_data.items():
    lrec = llm.get(bp, {})
    np = lrec.get("llm_named_pathogen", "").strip()
    if np:
        d["named_pathogen"].add(np)
    nh = lrec.get("llm_named_host", "").strip()
    if nh:
        d["named_host"].add(nh)

# ── Categorise ────────────────────────────────────────────────────────────────

categories = defaultdict(list)
for bp, d in bp_data.items():
    s = llm_setting(bp)
    t = llm_treat(bp)
    categories[(s, t)].append(d)

for key in categories:
    categories[key].sort(key=lambda d: d["n_coinf"], reverse=True)

# ── Document helpers ───────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_field(para, label, value):
    """Append 'Label: value' as bold+normal runs to an existing paragraph."""
    run = para.add_run(f"{label}: ")
    run.bold = True
    para.add_run(value or "—")

def add_divider(doc):
    """Thin horizontal rule between BioProjects."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

def coinf_summary(d):
    n   = d["n_biosamples"]
    nc  = d["n_coinf"]
    nhc = d["n_hc"]
    pct = f"{100*nc/n:.0f}%" if n else "0%"
    return f"{nc}/{n} BioSamples co-infected ({pct}); {nhc} HC (diff-genus)"

def pathogens_str(d):
    all_p = d["pathogens"]
    np    = d["named_pathogen"]
    if not all_p and not np:
        return "—"
    parts = []
    if np:
        parts.append(f"declared: {', '.join(sorted(np))}")
    if all_p:
        parts.append(f"STAT detected: {', '.join(sorted(all_p)[:8])}" +
                     (f" (+{len(all_p)-8} more)" if len(all_p) > 8 else ""))
    return "; ".join(parts)

def publication_str(d):
    pmid  = d["primary_pmid"]
    pub   = d["primary_publication"]
    date  = d["primary_pub_date"]
    sub   = d["bp_submission_date"]
    parts = []
    if pmid:
        parts.append(f"PMID {pmid}")
    if pub:
        parts.append(pub[:120] + ("…" if len(pub) > 120 else ""))
    if date:
        parts.append(f"({date[:7]})")
    elif sub:
        parts.append(f"submitted {sub[:7]}")
    return "  ".join(parts) if parts else "No PMID found"

def abstract_snippet(d, max_chars=300):
    text = d.get("abstract", "").strip()
    if not text:
        return ""
    return text[:max_chars] + ("…" if len(text) > max_chars else "")

# ── Write one document per category ───────────────────────────────────────────

for (setting, treat) in [(s, t) for s in ["field", "lab"] for t in TREAT_ORDER]:
    entries = categories.get((setting, treat), [])
    if not entries:
        continue

    fname   = f"{setting}_{treat}.docx"
    setting_label = "Field" if setting == "field" else "Lab / unclear / mixed"
    treat_label   = TREAT_LABELS[treat]

    doc = Document()
    set_narrow_margins(doc)

    # Title
    h = doc.add_heading(f"{setting_label}  ×  {treat_label}", level=0)
    h.runs[0].font.size = Pt(16)

    sub = doc.add_paragraph(
        f"{len(entries)} BioProjects   "
        f"(LLM setting: {setting}, treatment: {treat})"
    )
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()  # spacer

    for i, d in enumerate(entries, 1):
        bp   = d["BioProject"]
        lrec = llm.get(bp, {})

        # BioProject accession as sub-heading
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(6)
        hdr.paragraph_format.space_after  = Pt(2)
        run = hdr.add_run(f"{i}.  {bp}")
        run.bold = True
        run.font.size = Pt(11)
        run2 = hdr.add_run(
            f"   [{lrec.get('llm_study_setting','')} / {lrec.get('llm_treatment','')} "
            f"| confidence: {lrec.get('llm_confidence','')}]"
        )
        run2.font.size  = Pt(9)
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        # Title
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_field(p, "Title", d["title"] or "—")

        # Publication
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_field(p, "Publication", publication_str(d))

        # Hosts
        hosts_str = ", ".join(sorted(d["named_host"]))[:160] if d["named_host"] else "—"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_field(p, "Host(s)", hosts_str)

        # Pathogens
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_field(p, "Pathogen(s)", pathogens_str(d))

        # Co-infection summary
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_field(p, "Co-infection", coinf_summary(d))

        # LLM rationale
        rationale = lrec.get("llm_rationale", "").strip()
        if rationale:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            add_field(p, "LLM rationale", rationale[:400] + ("…" if len(rationale) > 400 else ""))

        # Abstract snippet
        snip = abstract_snippet(d)
        if snip:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_field(p, "Abstract", snip)

        # Notes line (blank, for handwriting)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("Notes: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        p.add_run("_" * 90)

        add_divider(doc)

    out_path = OUT_DIR / fname
    doc.save(str(out_path))
    print(f"Written: {out_path}  ({len(entries)} BioProjects)")

print(f"\nAll files in: {OUT_DIR}/")
